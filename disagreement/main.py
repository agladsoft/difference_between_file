import difflib
from docx import Document


def list_from_file(file_name):
    document = Document(file_name)
    paragraphs = []
    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)
    return paragraphs


def get_diff(file_name1, file_name2):
    diff_items = difflib.ndiff(list_from_file(file_name1), list_from_file(file_name2))
    diffs = []
    current_flag = '+'
    first_column = ' '
    second_column = ' '

    for diff in diff_items:
        '''if diff[0] == ' ':
            continue
        print(diff)'''

        if diff[0] == '-':
            if current_flag == '-':
                diffs.append((first_column, second_column))
            first_column = diff[2:]
            current_flag = '-'
        elif diff[0] == '+':
            second_column = diff[2:]
            diffs.append((first_column, second_column))
            first_column = second_column = ' '
            current_flag = '+'

    return diffs


def save_disagreement(name1, name2):
    rows = 1
    cols = 2
    result = Document()
    result.add_heading('Протокол разногласий')
    table = result.add_table(rows=rows, cols=cols)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Редакция заказчика'
    heading_cells[1].text = 'Редакция исполнителя'

    diffs = get_diff(name1, name2)
    for diff in diffs:
        cells = table.add_row().cells
        cells[0].text = diff[0]
        cells[1].text = diff[1]

    result.save('disagreement.docx')


if __name__ == '__main__':
    name1 = 'test1.docx'
    name2 = 'test2.docx'
    save_disagreement(name1, name2)
