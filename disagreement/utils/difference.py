import difflib
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
import io
import re

from disagreement.acceptable import skips, replacements


def list_from_file(document: str) -> list:
    paragraphs = [re.sub(r'\ {2,}', ' ', paragraph).replace('\xa0', ' ').strip() for paragraph in document.split('\n')]
    return paragraphs


def get_diff(list1: list, list2: list) -> list:
    diff_items = difflib.ndiff(list1, list2)
    diffs = []
    current_flag = None
    first_column = ''
    second_column = ''

    for diff in diff_items:
        if diff[0] == '-':
            if current_flag:
                diffs.append((first_column, second_column))
                second_column = ''
            first_column = diff[2:]
            current_flag = '-'
        elif diff[0] == '+':
            if current_flag == '+':
                diffs.append((first_column, second_column))
                first_column = ''
            second_column = diff[2:]
            current_flag = '+'

    return diffs


def save_disagreement(file1: str, file2: str, count_error: int) -> io.BytesIO:
    result = Document()
    result.add_heading('Протокол разногласий')
    table = result.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    table.autofit = False
    heading_cells = table.rows[0].cells
    heading_cells[0].text = '№'
    heading_cells[0].width = Inches(0.6)
    heading_cells[1].text = 'Редакция заказчика'
    heading_cells[1].width = Inches(3)
    heading_cells[2].text = 'Редакция исполнителя'
    heading_cells[2].width = Inches(3)

    list1, list2 = list_from_file(file1), list_from_file(file2)
    diffs = get_diff(list1, list2)
    for diff in diffs:
        match_number = re.search(r'^(\.?,?\d{0,2}){0,4} ', diff[0])
        number = match_number[0] if match_number else ''
        text1, text2 = [re.sub(r'(?:^(\.?,?\d{0,2}){0,4} |\.?,?$)', '', text).strip() for text in diff]

        cells = table.add_row().cells
        cells[0].width = Inches(0.6)
        cells[1].width = Inches(3)
        cells[2].width = Inches(3)
        cells[0].text = number

        left_diff_count = 0
        right_diff_count = 0

        sequence = difflib.SequenceMatcher(a=text1.lower(), b=text2.lower(), autojunk=False)

        paragraph1 = cells[1].paragraphs[0]
        paragraph2 = cells[2].paragraphs[0]

        for op, i1, i2, j1, j2 in sequence.get_opcodes():
            run1 = paragraph1.add_run(text1[i1:i2])
            run2 = paragraph2.add_run(text2[j1:j2])
            if op in ("delete", "insert"):
                if text1[i1:i2] not in skips:
                    run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    left_diff_count += i2 - i1
                if text2[j1:j2] not in skips:
                    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    right_diff_count += j2 - j1
            if op == "replace":
                if (text1[i1:i2], text2[j1:j2]) not in replacements and (text2[j1:j2], text1[i1:i2]) not in replacements:
                    run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    left_diff_count += i2 - i1
                    right_diff_count += j2 - j1

        if left_diff_count <= count_error and right_diff_count <= count_error:
            table._tbl.remove(table.rows[-1]._tr)

    file_stream = io.BytesIO()
    result.save(file_stream)
    file_stream.seek(0)

    return file_stream


